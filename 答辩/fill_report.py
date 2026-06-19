# -*- coding: utf-8 -*-
"""填写期末大作业报告模板"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import copy

SRC = r'C:\Users\Administrator\Desktop\名字+学号+班级 python程序设计 期末大作业模版.docx'
DST = r'd:\python期末作业\贵州旅游景点问答助手\答辩\贵州旅游景点智能问答助手_课程报告.docx'

doc = Document(SRC)

def replace_para_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def find_para(doc, text):
    for para in doc.paragraphs:
        if text in para.text:
            return para
    return None

def remove_by_texts(doc, texts):
    to_rm = []
    for p in doc.paragraphs:
        for t in texts:
            if t in p.text:
                to_rm.append(p)
                break
    for p in to_rm:
        el = p._element
        el.getparent().remove(el)

# === 删除指导性段落 ===
remove_by_texts(doc, [
    '内容要求', '格式要求', '本章用简洁清楚', '对应评分项', '本节不单独给分',
    '这是文字报告', '本章是文字报告', '重点看数据表', '小四号宋体', '首行缩进2字符',
    '行距1.5倍', '一级标题：四号黑体', '二级标题：小四号黑体', '本章字数建议',
    '表格内文字用小四', '目录结构可用代码块', '标题格式同上', '图形居中、清晰',
    '按"图 1', '每张截图配一句说明', '正文：小四号宋体', '关键词之间用空格隔开',
    '结尾无句号', '给出5到10个关键词', '关键词应主要覆盖', '核心技术：',
    '应用领域：', '技术特性：', '数据库与建模：', '字数：约300字', '说明：为何做',
    '绘制系统的 E-R 图', '把图复制粘贴到此处', '简要说明表设计依据',
    '数据表如无正当理由', '字段重复严重', '一张表塞下', '同一信息在多张',
    '该用 user_id', '字段命名无法表达', '归纳总结本项目', '小结你对核心功能',
    '说明项目的不足', '说明 AI 工具', '指出后续可优化', '本章字数建议 400',
    '功能名称：要分析', '操作流程：用户在页面', '前端字段：前端收集',
    '后端处理：后端接收', '数据库存储：哪些字段', '返回结果：后端返回',
    '页面变化：前端拿到', '按上面七个要素', '用与 4.1 相同的方式',
    '首页 / 主界面截图', '登录或注册页面截图', '核心功能页面截图',
    '操作结果截图', '数据库表数据截图', '数据流可参考下列写法',
    '数据流可用文字箭头', '请替换为本人项目的',
])

# === 摘要 ===
p = find_para(doc, '字数：约300字')
if p:
    replace_para_text(p,
        '随着旅游业的快速发展，游客对旅游信息的需求日益增长。然而，现有的旅游信息散落在'
        '多个平台上，游客需要在携程、马蜂窝、小红书等多个App之间来回切换才能获取完整的'
        '景点信息，查找效率低下。针对这一问题，本项目设计并实现了一个基于Python Flask的'
        '贵州旅游景点智能问答助手。系统采用RAG（检索增强生成）技术，通过TF-IDF向量检索'
        '和意图识别，实现了自然语言提问、秒级响应的智能问答功能。系统后端使用Python Flask'
        '搭建RESTful API，前端使用HTML、CSS和JavaScript实现交互界面，数据库使用SQLite'
        '存储景点、图片、攻略、路线等数据。项目共设计了6张数据表，覆盖117个贵州景点数据，'
        '实现了智能问答、景点地图可视化、信息管理等核心功能，并通过Leaflet地图实现了117个'
        '景点的交互式标注和路线可视化。'
    )

# === 关键词 ===
for p in doc.paragraphs:
    if '关键词：' in p.text and '内容' not in p.text:
        replace_para_text(p,
            '关键词：Python Flask TF-IDF RAG 智能问答 旅游信息检索 意图识别 '
            'SQLite RESTful API 前后端分离 Leaflet 地图可视化 n-gram分词'
        )
        break

# === 第一章 ===
p = find_para(doc, '第一章')
if p:
    replace_para_text(p,
        '1.1 项目背景\n\n'
        '贵州省拥有丰富的旅游资源，包括黄果树瀑布、西江千户苗寨、荔波小七孔等众多知名景点。'
        '然而，游客在规划贵州旅行时，面临着旅游信息碎片化的问题——景点的门票价格、开放时间、'
        '交通方式、游玩攻略等信息散落在携程、马蜂窝、小红书等多个平台上，游客需要在多个App'
        '之间来回切换，翻阅十几个页面才能拼凑出完整的旅行计划。\n\n'
        '1.2 使用对象\n\n'
        '本系统面向计划前往贵州旅游的游客，包括学生、家庭游客、自助旅行者等群体。同时，'
        '系统也为旅游从业者提供景点信息管理功能。\n\n'
        '1.3 用户痛点\n\n'
        '用户在查找旅游信息时面临三个核心痛点：第一，信息碎片化，同一景点的信息分散在多个'
        '平台；第二，查找效率低，获取一个景点的完整信息需要在多个App之间切换；第三，信息'
        '不精准，搜索引擎返回的是链接列表而非直接答案。\n\n'
        '1.4 解决方案\n\n'
        '本系统通过RAG检索增强生成技术，将117个贵州景点的数据整合到统一的数据库中，用户'
        '只需用自然语言输入问题，系统即可在毫秒级时间内返回精准答案，并附带来源信息。核心'
        '功能包括：智能问答（意图识别+语义检索）、景点地图（117个景点可视化标注）、信息'
        '管理（景点/攻略/路线的增删改查）。'
    )

# === 2.1 系统功能概述 ===
p = find_para(doc, '列出系统的主要功能模块')
if p:
    replace_para_text(p,
        '本系统包含以下核心功能模块：（1）智能问答模块：接收用户自然语言提问，通过意图识别'
        '判断问题类型，利用TF-IDF向量检索匹配相关景点，生成带来源的精准回答。（2）景点地图'
        '模块：基于Leaflet地图标注117个贵州景点，支持分类筛选、路线可视化和景点详情查看。'
        '（3）信息管理模块（管理后台）：对景点、攻略、路线数据进行增删改查操作。（4）数据'
        '可视化模块：以图表形式展示景点分布、类别统计等数据。（5）收藏功能：用户可收藏'
        '感兴趣的景点。'
    )

# === 2.2 用户角色 ===
p = find_para(doc, '描述系统中的用户角色')
if p:
    replace_para_text(p,
        '本系统设置两种用户角色：（1）普通用户：可使用智能问答、浏览景点地图、查看景点详情、'
        '收藏景点、查看路线推荐和数据可视化大屏。（2）管理员：除普通用户的所有权限外，还可'
        '访问管理后台，对景点、攻略、路线数据进行新增、编辑、删除操作。'
    )

# === 2.3 技术栈表格 ===
t = doc.tables[0]
# 先更新已有行
data = [
    ['前端', 'HTML + CSS + JavaScript', '页面结构、样式与交互逻辑'],
    ['地图', 'Leaflet.js', '交互式地图标注与路线可视化'],
    ['后端', 'Python + Flask', '提供RESTful API、处理业务逻辑'],
    ['数据库', 'SQLite', '存储景点、图片、攻略、路线等数据'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        t.rows[i+1].cells[j].text = val
# 更新最后一行
t.rows[4].cells[0].text = 'AI核心'
t.rows[4].cells[1].text = 'TF-IDF + 意图识别 + RAG'
t.rows[4].cells[2].text = '语义检索、意图分类、回答生成'

# === 2.4 目录结构 ===
p = find_para(doc, '用下面的目录树说明')
if p:
    replace_para_text(p,
        '项目目录结构如下：\n\n'
        '贵州旅游景点问答助手/\n'
        '  app/                    # Flask应用核心\n'
        '    main.py              # 应用入口，注册蓝图\n'
        '    database.py          # 数据库建表与数据导入\n'
        '    db.py                # 数据库连接工具\n'
        '    admin.py             # 管理后台蓝图\n'
        '    routes/              # 路由模块\n'
        '      pages.py           # 页面路由\n'
        '      attractions.py     # 景点API\n'
        '      route_api.py       # 路线API\n'
        '      favorites.py       # 收藏API\n'
        '      stats.py           # 统计API\n'
        '    qa/                  # 问答系统模块\n'
        '      __init__.py        # RAG问答门面类\n'
        '      search.py          # TF-IDF向量检索\n'
        '      intent.py          # 意图识别\n'
        '      generator.py       # 回答生成\n'
        '  static/                # 前端静态资源\n'
        '    css/                 # 样式文件\n'
        '    js/                  # JavaScript文件\n'
        '    images/              # 图片资源\n'
        '  templates/             # HTML模板\n'
        '  data/                  # 数据文件\n'
        '    guizhou_travel.db    # SQLite数据库\n'
        '    guizhou_attractions.json  # 景点JSON数据\n'
        '  scripts/               # 数据爬虫脚本\n'
        '  run.py                 # 启动脚本\n'
        '  requirements.txt       # 依赖清单'
    )

# === 3.1 数据表 ===
p = find_para(doc, '列出系统的主要数据表')
if p:
    replace_para_text(p, '本系统共设计6张数据表，各表作用和主要字段如下表所示。')

t2 = doc.tables[2]
td = [
    ['attractions', '存储景点基本信息', 'id, name, address, category, ticket_price, opening_hours, description, rating, latitude, longitude'],
    ['images', '存储景点图片', 'id, attraction_id(FK), image_url, image_type, description, is_primary'],
    ['guides/routes/favorites', '攻略/路线/收藏', '外键关联attractions表'],
]
for i, row in enumerate(td):
    for j, val in enumerate(row):
        t2.rows[i+1].cells[j].text = val

# === 3.2 E-R图 ===
p = find_para(doc, '绘制系统的 E-R 图')
if p:
    replace_para_text(p,
        '本系统的表关系如下：attractions表为主表，images表和guides表通过attraction_id外键'
        '关联attractions表，形成一对多关系（一个景点可有多张图片和多篇攻略）。favorites表'
        '通过attraction_id关联attractions表，记录用户的收藏关系。routes表为独立表，通过'
        'attractions_list字段以文本形式记录路线包含的景点名称。各表之间关系清晰，通过外键'
        '约束保证数据一致性。'
    )

# === 3.3 三大范式 ===
p = find_para(doc, '简要说明表设计依据')
if p:
    replace_para_text(p,
        '本系统的数据库设计遵循数据库三大范式：第一范式（1NF）：每张表的每个字段都是不可再分'
        '的原子数据项，如attractions表中的ticket_price、opening_hours等字段均存储单一值。'
        '第二范式（2NF）：在满足第一范式的基础上，非主键字段完全依赖于主键。如images表中的'
        'image_url、image_type等字段都直接依赖于images表的主键id。第三范式（3NF）：在满足'
        '第二范式的基础上，非主键字段之间不存在传递依赖。如attractions表中，景点的名称、地址、'
        '类别等属性都直接依赖于主键id。routes表中的attractions_list字段以文本形式存储路线'
        '包含的景点名称，存在一定的数据冗余，但考虑到路线功能相对独立且查询频率较低，这种设计'
        '在项目规模下是合理的。'
    )

# === 4.1 核心功能一 ===
p = find_para(doc, '4.1')
if p:
    replace_para_text(p,
        '4.1 核心功能一：智能问答\n\n'
        '功能名称：智能问答（RAG检索增强生成）\n\n'
        '操作流程：用户在首页搜索框或聊天界面输入自然语言问题，如[黄果树瀑布门票多少钱]，'
        '点击发送按钮。\n\n'
        '前端字段：前端通过POST请求发送JSON数据 {message: 用户输入的问题文本}。\n\n'
        '后端处理：后端接收请求后，调用RAG问答系统。首先通过n-gram分词对问题进行分词，'
        '然后通过意图识别模块判断问题类型（门票/交通/时间/推荐/综合），接着使用TF-IDF'
        '向量检索在117个景点中找到最相关的Top-5景点，最后将检索结果和问题交给回答生成'
        '模块，生成带来源的完整回答。\n\n'
        '数据库存储：问答过程为只读操作，不写入数据库。系统从attractions表读取景点信息'
        '用于检索和回答生成。\n\n'
        '返回结果：后端返回JSON数据，包含answer（回答文本）、attractions（推荐景点列表，'
        '含id、name、rating、ticket_price等）、suggestions（追问建议列表）。\n\n'
        '页面变化：前端ChatSystem组件接收到响应后，以流式效果逐字显示回答文本，并在回答'
        '下方展示景点推荐卡片和追问建议按钮。'
    )

# === 字段说明表 ===
t3 = doc.tables[3]
fd = [
    ['message', '用户输入的问题', '前端输入框', '后端qa.answer_question()', '是'],
    ['answer', '生成的回答文本', '后端生成', '前端流式显示', '是'],
    ['attractions', '推荐景点列表', '数据库检索', '前端卡片展示', '是'],
    ['suggestions', '追问建议', '后端生成', '前端按钮展示', '否'],
]
for i, row in enumerate(fd):
    for j, val in enumerate(row):
        t3.rows[i+1].cells[j].text = val

# === 4.2 核心功能二 ===
p = find_para(doc, '4.2')
if p:
    replace_para_text(p,
        '4.2 核心功能二：景点地图可视化\n\n'
        '功能名称：景点地图可视化\n\n'
        '操作流程：用户点击导航栏[景点地图]进入地图页面，地图自动加载并标注117个贵州景点。'
        '用户可通过侧边栏按类别筛选景点，点击景点标记可查看详情。\n\n'
        '前端字段：前端通过GET请求调用/api/attractions接口获取景点列表，包含id、name、'
        'address、category、latitude、longitude等字段。\n\n'
        '后端处理：后端从attractions表查询所有景点数据，按类别参数过滤（如有），返回JSON'
        '格式的景点列表。\n\n'
        '数据库存储：只读操作，从attractions表读取景点坐标和基本信息。\n\n'
        '返回结果：返回景点列表JSON，每个景点包含id、name、address、category、latitude、'
        'longitude、ticket_price、rating等字段。\n\n'
        '页面变化：Leaflet地图根据返回的经纬度数据在地图上标注景点标记，侧边栏显示景点'
        '列表，支持按类别筛选和路线可视化。'
    )

# === 4.3 核心功能三 ===
p = find_para(doc, '4.3')
if p:
    replace_para_text(p,
        '4.3 核心功能三：景点信息管理（CRUD）\n\n'
        '功能名称：景点信息管理\n\n'
        '操作流程：管理员登录管理后台，在景点管理页面可查看景点列表，点击[新增]按钮添加'
        '新景点，点击[编辑]修改景点信息，点击[删除]按钮删除景点。\n\n'
        '前端字段：新增/编辑时前端收集name、address、category、ticket_price、opening_hours、'
        'description、features、rating等字段，通过POST/PUT请求提交。\n\n'
        '后端处理：后端接收数据后进行非空校验，然后执行INSERT或UPDATE SQL语句操作attractions'
        '表。删除操作执行DELETE语句，并级联清理关联的images和guides记录。\n\n'
        '数据库存储：写入attractions表，新增景点时系统自动生成id和created_at。编辑时更新'
        '对应记录的所有字段。\n\n'
        '返回结果：操作成功返回{success: true, message: 操作成功}，失败返回{error: 错误信息}。\n\n'
        '页面变化：操作成功后自动刷新景点列表，新增的景点出现在列表中，删除的景点从列表中'
        '移除。'
    )

# === 第五章 截图说明 ===
p = find_para(doc, '用截图说明你确实做了这些功能')
if p:
    replace_para_text(p,
        '本章展示系统各核心功能的运行截图，证明功能已实现。\n\n'
        '注：截图请在运行项目后自行截取并替换以下占位说明。\n\n'
        '图1 首页界面：展示系统首页的全屏Banner、搜索框和快捷标签。\n\n'
        '图2 智能问答：用户输入问题后，系统返回带来源的回答和景点推荐卡片。\n\n'
        '图3 景点地图：Leaflet地图标注117个贵州景点，支持分类筛选。\n\n'
        '图4 景点详情：展示景点的图片、门票、交通、攻略等详细信息。\n\n'
        '图5 管理后台：管理员可对景点、路线进行增删改查操作。\n\n'
        '图6 数据可视化：以图表形式展示景点分布和类别统计。'
    )

# === 第六章 总结 ===
p = find_para(doc, '归纳总结本项目已完成的主要功能')
if p:
    replace_para_text(p,
        '6.1 项目总结\n\n'
        '本项目完成了贵州旅游景点智能问答助手的设计与实现，主要成果包括：（1）完成了117个'
        '贵州景点数据的采集、清洗和入库，设计了6张数据表。（2）实现了基于TF-IDF和RAG的'
        '智能问答系统，支持5种意图分类，响应时间在毫秒级。（3）实现了Leaflet交互式地图，'
        '支持117个景点的标注、筛选和路线可视化。（4）完成了管理后台，支持景点、攻略、路线'
        '的增删改查操作。（5）实现了数据可视化大屏、收藏功能和响应式UI设计。\n\n'
        '对核心功能数据流的理解：智能问答功能的数据流为[用户输入->意图识别->TF-IDF检索->'
        '景点匹配->回答生成->结果返回]，每个环节都有明确的输入输出和字段传递。景点地图的'
        '数据流为[前端请求->后端查询->JSON返回->地图标注]，字段来源清晰、去向明确。\n\n'
        '6.2 不足与改进\n\n'
        '项目目前存在以下不足：（1）缺少用户账号系统，无法实现个性化推荐和多用户隔离。'
        '（2）不支持多轮对话记忆，每次问答都是独立的，无法理解上下文。（3）景点数据来源'
        '有限，部分景点信息不够完整。（4）未实现语音交互功能。\n\n'
        '6.3 AI工具使用说明\n\n'
        '本项目在开发过程中使用了AI工具辅助：使用AI辅助编写部分代码逻辑、生成文档模板、'
        '优化代码结构。但所有核心算法（TF-IDF检索、意图识别、回答生成）的理解和实现均为'
        '本人完成。\n\n'
        '6.4 后续优化方向\n\n'
        '（1）引入用户账号系统，支持登录注册和个人中心。（2）接入大语言模型实现多轮对话'
        '记忆。（3）持续爬取更新景点数据，丰富信息覆盖面。（4）开发移动端适配或小程序版本。'
    )

# === E-R图占位 ===
p = find_para(doc, '请替换为本人项目的 E-R 图')
if p:
    replace_para_text(p, '（E-R图见上文3.2节表关系说明）')

# === 数据流示例图 ===
p = find_para(doc, '核心功能一数据流（示例）')
if p:
    replace_para_text(p, '智能问答数据流：用户输入 -> 意图识别 -> TF-IDF检索 -> 景点匹配 -> 回答生成 -> 结果返回')

# === 保存 ===
doc.save(DST)
print('报告已生成:', DST)
